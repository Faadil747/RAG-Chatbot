"""DOCX text extraction using python-docx."""
from __future__ import annotations

import io
import logging

logger = logging.getLogger("ai-service.parsing.docx")


def extract_docx_text(file_bytes: bytes) -> str:
    from docx import Document

    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception:
        logger.exception("python-docx failed to open the document")
        return ""

    chunks: list[str] = []

    for para in doc.paragraphs:
        if para.text.strip():
            chunks.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                chunks.append(" | ".join(cells))

    return "\n".join(chunks).strip()
